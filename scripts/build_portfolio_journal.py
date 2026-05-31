from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "Reyniel_Polancos_Portfolio_Journal.docx"
PROFILE_IMAGE = ROOT / "images" / "profile" / "reyniel-polancos.png"
MOBILELEX_IMAGE = ROOT / "images" / "projects" / "mobilelex-dashboard.jpg"

PURPLE = "5E1B7A"
MAGENTA = "C95CFF"
DEEP = "1A052A"
LIGHT = "F7ECFF"
MUTED = "6B5878"
PANEL = "F6EFFA"
LINE = "D9BEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_w = cell._tc.get_or_add_tcPr().tcW
            tc_w.type = "dxa"
            tc_w.w = widths[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_text(paragraph, text, bold=False, italic=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("2C2233")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 24, DEEP, 0, 8),
        ("Heading 1", 16, PURPLE, 12, 6),
        ("Heading 2", 12.5, PURPLE, 8, 4),
        ("Heading 3", 11, DEEP, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.style = doc.styles["Normal"]
        add_text(footer, "Reyniel Polancos | Fresh IT Graduate Portfolio", size=8.5, color=MUTED)


def add_cover(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [5900, 3300])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_border(cell, "FFFFFF", "0")
        set_cell_margins(cell, 120, 120, 120, 120)

    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    add_text(p, "PORTFOLIO JOURNAL", bold=True, size=9.5, color=MAGENTA)
    title = left.add_paragraph()
    title.style = doc.styles["Title"]
    add_text(title, "Reyniel Polancos", bold=True, size=28, color=DEEP)
    subtitle = left.add_paragraph()
    add_text(subtitle, "Fresh IT Graduate | Frontend & Mobile App Development", bold=True, size=12, color=PURPLE)
    objective = left.add_paragraph()
    add_text(
        objective,
        "A concise portfolio prepared for journal requirements, highlighting academic background, technical skills, projects, certifications, and career readiness for entry-level developer roles.",
        size=10.5,
        color="3C2D46",
    )

    facts = [
        ("Current Status", "Fresh Graduate"),
        ("Target Role", "Junior Developer"),
        ("Focus", "Frontend & Mobile App"),
        ("Location", "Digos City, Davao del Sur"),
        ("Email", "reynren11@gmail.com"),
        ("GitHub", "github.com/LuckyYunqi"),
    ]
    fact_table = left.add_table(rows=len(facts), cols=2)
    set_table_width(fact_table, [1950, 3650])
    for row, (label, value) in zip(fact_table.rows, facts):
        for cell in row.cells:
            set_cell_border(cell, LINE, "6")
            set_cell_margins(cell, 90, 110, 90, 110)
        set_cell_shading(row.cells[0], PANEL)
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(PURPLE)

    if PROFILE_IMAGE.exists():
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.add_run().add_picture(str(PROFILE_IMAGE), width=Inches(2.2))
    cap = right.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap, "Prepared for Portfolio / Journal Submission", bold=True, size=10.5, color=PURPLE)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9000])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, PANEL)
    set_cell_border(cell, LINE, "8")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    add_text(p, title, bold=True, size=11, color=PURPLE)
    p2 = cell.add_paragraph()
    add_text(p2, body, size=10.3, color="3A2B43")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text)


def add_sections(doc):
    doc.add_heading("Career Objective", level=1)
    add_callout(
        doc,
        "Objective",
        "To contribute as an entry-level frontend or mobile app developer while strengthening real-world development practices, code quality, collaboration, and user-centered problem solving.",
    )

    doc.add_heading("About Me", level=1)
    p = doc.add_paragraph()
    add_text(
        p,
        "I am Reyniel Polancos, a fresh IT graduate with a foundation in web design, frontend development, and practical React-based project work. I build clean, responsive interfaces and continue improving through hands-on projects using HTML, CSS, React JS, and React Native.",
    )
    p = doc.add_paragraph()
    add_text(
        p,
        "I have beginner-level practical experience in mobile app development, including React Native app structure, screens, navigation concepts, and connecting user-focused features into a usable interface.",
    )

    doc.add_heading("Skills Profile", level=1)
    skills = [
        ("Frontend", "HTML, CSS, React JS"),
        ("Mobile", "React Native"),
        ("Design / Editing", "Photoshop, Premiere Pro, After Effects"),
        ("Tools", "GitHub, VS Code, XAMPP"),
    ]
    table = doc.add_table(rows=1 + len(skills), cols=2)
    set_table_width(table, [2500, 6500])
    headers = ("Category", "Skills")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PURPLE)
        set_cell_border(cell, PURPLE)
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
    for row, (category, items) in zip(table.rows[1:], skills):
        row.cells[0].text = category
        row.cells[1].text = items
        for cell in row.cells:
            set_cell_border(cell, LINE, "6")
        set_cell_shading(row.cells[0], PANEL)

    doc.add_heading("Education", level=1)
    education = [
        ("University of Mindanao", "College Graduate", "Digos City, Davao del Sur"),
        ("University of Mindanao", "Senior High School", "Digos City, Davao del Sur"),
        ("Holy Cross Academy of Digos", "Junior High School", "Digos City, Davao del Sur"),
        ("Ramon Magsaysay Elementary School", "Elementary", "Digos City, Davao del Sur"),
    ]
    for school, level, location in education:
        p = doc.add_paragraph()
        add_text(p, school, bold=True, color=DEEP)
        add_text(p, f" - {level} | {location}", color=MUTED)

    doc.add_page_break()
    doc.add_heading("Projects", level=1)
    doc.add_heading("Lawyer Consultation Mobile App", level=2)
    add_bullet(doc, "Type: Mobile App")
    add_bullet(doc, "Role: Mobile App Developer")
    add_bullet(doc, "Technology: React Native, mobile app screens, consultation workflows")
    add_bullet(doc, "Repository: github.com/xxexil/MobileLex")
    p = doc.add_paragraph()
    add_text(p, "Description: ", bold=True)
    add_text(
        p,
        "A lawyer consultation app focused on schedules, messages, payments, and client-centered consultation flows.",
    )
    if MOBILELEX_IMAGE.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(MOBILELEX_IMAGE), width=Inches(2.25))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(caption, "MobileLex dashboard screen", italic=True, size=9, color=MUTED)

    projects = [
        ("Inventory of Medicines System (IMES)", "Full Stack / System Contributor", "Web app, inventory management, reports", "A web-based system for managing pharmaceutical inventory, stock movements, and reporting. Screenshot unavailable because the system was turned over to the client."),
        ("Personal Portfolio Website", "Web Designer / Frontend Developer", "HTML, CSS, JavaScript", "A responsive portfolio website presenting profile, education, certificates, projects, skills, and contact information."),
        ("Responsive Web Design", "Frontend Practice", "HTML, CSS, accessibility, responsive layouts", "Coursework and practice projects focused on responsive layouts and mobile-first fundamentals."),
    ]
    for title, role, stack, description in projects:
        doc.add_heading(title, level=2)
        add_bullet(doc, f"Role: {role}")
        add_bullet(doc, f"Tech Stack: {stack}")
        add_bullet(doc, description)

    doc.add_page_break()
    doc.add_heading("Certifications", level=1)
    certs = [
        ("Installing and Configuring Computer Systems", "TESDA", "Completed September 5, 2025"),
        ("Setting Up Computer Networks", "TESDA", "Completed September 5, 2025"),
        ("Setting Up Computer Servers", "TESDA", "Completed September 5, 2025"),
        ("Maintaining Computer Systems and Networks", "TESDA", "Completed September 5, 2025"),
        ("Introduction to CSS", "TESDA", "Completed September 1, 2025"),
        ("IT Specialist: HTML and CSS", "Certiport / Pearson VUE", "Completed February 28, 2026"),
    ]
    for title, issuer, date in certs:
        add_bullet(doc, f"{title} - {issuer} ({date})")

    doc.add_heading("OJT / Career Preparation", level=1)
    add_callout(
        doc,
        "Readiness Statement",
        "I am prepared to enter a real development environment as a fresh graduate who is willing to learn, accept feedback, document tasks, communicate clearly, and improve through daily practice.",
    )
    add_bullet(doc, "Strengthen React JS and React Native fundamentals through small, complete projects.")
    add_bullet(doc, "Practice GitHub workflow, code organization, and project documentation.")
    add_bullet(doc, "Continue improving UI consistency, responsive layouts, and accessibility.")
    add_bullet(doc, "Prepare for interviews by explaining project role, decisions, challenges, and learning outcomes.")

    doc.add_heading("Reflection", level=1)
    p = doc.add_paragraph()
    add_text(
        p,
        "This portfolio represents my current growth as a fresh IT graduate. It shows my academic foundation, beginner-level practical experience, and willingness to continue learning in a company related to my expertise.",
    )


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    doc.add_page_break()
    add_sections(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
